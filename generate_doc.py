"""Generate the QA Jira CLI project documentation as a Word .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT_PATH = "/Users/salescode/Downloads/QA_Jira_CLI_Documentation.docx"
TODAY = datetime.date(2026, 6, 1).strftime("%B %d, %Y")

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
for margin in ("top_margin","bottom_margin","left_margin","right_margin"):
    setattr(section, margin, Inches(1))

# ── Style helpers ───────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(18)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(14)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(12)
    return p

def body(text):
    return doc.add_paragraph(text)

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A93")
        tcPr.append(shd)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        fill = "EBF5FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            tc = cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for ri2, row2 in enumerate(t.rows):
            for ci2, w in enumerate(col_widths):
                row2.cells[ci2].width = Inches(w)
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("QA Jira CLI (Python)")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x93)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Complete Technical Documentation")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f"Date: {TODAY}").font.size = Pt(12)

repo_p = doc.add_paragraph()
repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo_p.add_run("Repository: https://github.com/Hiyaarora/Qa-jira-py").font.size = Pt(11)

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.add_run("Author: Hiya Arora").font.size = Pt(11)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
h1("1. PROJECT OVERVIEW")

body(
    "QA Jira CLI is a Python command-line tool built for QA engineers. "
    "It replaces the manual process of writing Jira bug reports and task "
    "descriptions by using an AI language model to structure plain-English "
    "input into professional, detailed Jira tickets. It can also export all "
    "bugs under a Jira epic into a styled Excel file. Five commands are "
    "available, all driven by an interactive arrow-key menu."
)
doc.add_paragraph()

body("Key capabilities:")
for cap in [
    "AI-structured bug reports — type a sentence, get 5 steps + actual/expected result",
    "AI-structured QA task descriptions — tested story / wrote test cases / other",
    "Vision AI — attach a screenshot and the AI reads it to write the description",
    "Auto-discovery of required Jira custom fields (e.g. Complexity) via createmeta API",
    "Excel bug-sheet export — all bugs in an epic → styled .xlsx with hyperlinks",
    "Interactive arrow-key menu — no command memorisation needed",
    "Resilient AI fallback — queries OpenRouter live for working free models",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(cap).font.size = Pt(10)

doc.add_paragraph()
h2("Tech Stack")
table(
    ["Layer", "Technology", "Why it was chosen"],
    [
        ["Language",        "Python 3.11+",         "Modern type hints, match statements, tomllib built-in"],
        ["Packaging",       "uv + pyproject.toml",  "Fast, reproducible installs; single tool for venv + deps"],
        ["HTTP client",     "httpx",                "Sync + async, clean API, replaces requests/aiohttp"],
        ["Interactive CLI", "questionary",          "Arrow-key menus and password prompts; closest match to @inquirer/prompts from the original Node.js CLI"],
        ["Terminal output", "rich",                 "Colours, spinners, dividers — replaces chalk"],
        ["Data models",     "pydantic v2",          "Validated type-safe models for Config, Issue, Bug, AttachmentInfo; field-level validation at the boundary"],
        ["Excel output",    "openpyxl",             "Full cell styling, hyperlinks, frozen rows — no external dependency on Google"],
        ["Date handling",   "python-dateutil",      "Parse any date format from Jira API responses"],
        ["Jira API",        "Jira REST API v3",     "Official cloud REST API; Basic-auth with email + API token"],
        ["AI provider",     "OpenRouter / Anthropic","OpenAI-compatible /chat/completions endpoint; no SDK dependency — raw httpx POST identical to the original Node.js CLI"],
        ["Testing",         "pytest + httpx.MockTransport", "Unit tests with mock Jira server; no real network calls in CI"],
        ["Config storage",  "~/.qa-jira/config.json","Same location as the original CLI; chmod 600; no env-var juggling"],
    ],
    col_widths=[1.3, 1.7, 3.5],
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. ENTRY POINTS
# ═══════════════════════════════════════════════════════════════════════════
h1("2. ENTRY POINTS")

h2("2.1  Global Command — jira")
body(
    "Installed by uv tool install (or uv pip install -e .) which registers the "
    "console_script defined in pyproject.toml:"
)
code('[project.scripts]\njira = "qa_jira.cli:main"')
body(
    "When you type jira in any terminal, the OS resolves /Users/salescode/.local/bin/jira "
    "which calls qa_jira.cli.main()."
)
doc.add_paragraph()

h2("2.2  cli.py  —  main()")
body("File: src/qa_jira/cli.py")
body("What happens when jira is invoked with no arguments (the normal case):")
for step in [
    "1.  sys.argv[1:] is empty → _interactive_menu() is called",
    "2.  questionary.select() renders an arrow-key menu of 6 choices",
    "3.  User picks a command → _run_command(cmd, rest) dispatches to the matching module",
    "4.  The module's run() function takes over",
]:
    body(step)
code(
    "def main() -> None:\n"
    "    args = sys.argv[1:]\n"
    "    if not args or args[0] in ('-h', '--help', 'help'):\n"
    "        _interactive_menu()\n"
    "        return\n"
    "    cmd, *rest = args\n"
    "    _run_command(cmd, rest)"
)
body(
    "When jira is invoked with arguments (e.g. jira mk bug) the menu is skipped "
    "and _run_command() is called directly — same effect."
)
doc.add_paragraph()

h2("2.3  Module: __main__.py  (python -m qa_jira)")
body("File: src/qa_jira/__main__.py")
body(
    "Allows running the package as a module. "
    "Imports and calls main() from cli.py — identical behaviour."
)
code("from qa_jira.cli import main\nif __name__ == '__main__':\n    main()")
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. FILE DEPENDENCY MAP
# ═══════════════════════════════════════════════════════════════════════════
h1("3. FILE DEPENDENCY MAP")
body(
    "Every source file in src/qa_jira/ is listed below. "
    "Arrows show what each file imports and which files import it."
)
doc.add_paragraph()
table(
    ["File", "Imports from", "Imported by", "Role (one line)"],
    [
        ["cli.py",                      "sys, questionary, rich\ncommands.*",                                  "pyproject.toml console_scripts",                        "Entry point; dispatches argv or shows arrow-key menu"],
        ["config.py",                   "json, os, sys, pathlib\nmodels.Config, rich",                         "commands.setup, commands.*, all run() functions",       "Load/save ~/.qa-jira/config.json with chmod 600"],
        ["models.py",                   "pydantic, typing",                                                    "config, jira_client, ai/*, commands/*",                 "All Pydantic data models for the project"],
        ["adf.py",                      "(stdlib only)",                                                       "jira_client, ai/base",                                  "Build Atlassian Document Format JSON for Jira descriptions"],
        ["file_handler.py",             "pathlib, dataclasses",                                                "commands/_helpers",                                     "Detect input type (file/google-sheet/url), validate file size"],
        ["jira_client.py",              "httpx, base64, re, pathlib\nadf, models",                            "commands/setup, commands/mk_bug, commands/task_create, commands/mk_bugsheet, commands/rm", "All Jira REST API calls — read, write, search, transition, delete"],
        ["prompts.py",                  "models (type hints only)",                                            "ai/__init__",                                           "System + user prompt templates for bug and task AI calls"],
        ["excel.py",                    "openpyxl, dateutil, pathlib\nmodels",                                 "commands/mk_bugsheet",                                  "Write 12-column .xlsx bugsheet with styling and hyperlinks"],
        ["ai/__init__.py",              "ai/base, ai/http_provider\nmodels, prompts",                         "commands/mk_bug, commands/task_create",                 "Public AI API — generate_bug_description(), generate_task_description()"],
        ["ai/base.py",                  "json, re, typing\nadf, models",                                      "ai/__init__",                                           "JSON parsing, ADF construction, AIBugResult/AITaskResult builders"],
        ["ai/http_provider.py",         "httpx, base64, pathlib\nmodels",                                     "ai/__init__",                                           "Raw httpx POST to /chat/completions; model fallback + vision support"],
        ["ai/anthropic_provider.py",    "anthropic SDK\nmodels",                                               "(unused — kept for Anthropic SDK path if needed)",      "Anthropic SDK-based provider (superseded by http_provider)"],
        ["ai/openai_compat_provider.py","openai SDK\nmodels",                                                  "(unused — kept for reference)",                         "OpenAI SDK-based provider (superseded by http_provider)"],
        ["commands/__init__.py",        "(empty)",                                                             "Python package marker",                                 "Package marker"],
        ["commands/_helpers.py",        "httpx, questionary, rich\njira_client, file_handler, models",        "commands/mk_bug, commands/task_create, commands/mk_bugsheet", "Shared interactive pickers: project, epic, user, attachment, required-fields"],
        ["commands/_attachment.py",     "httpx, rich\njira_client, models",                                   "commands/mk_bug, commands/task_create",                 "Upload file attachments; Google Sheet URLs go in ADF description"],
        ["commands/setup.py",           "httpx, questionary, rich, webbrowser\nconfig, jira_client, models",  "cli._run_command",                                      "4-step wizard: Jira creds + AI provider + key + validation"],
        ["commands/rm.py",              "httpx, questionary, rich\nconfig, jira_client",                      "cli._run_command",                                      "Fetch issue → display → confirm-No-by-default → delete"],
        ["commands/mk_bug.py",          "httpx, questionary, rich, sys\nconfig, jira_client, ai, models\ncommands/_helpers, commands/_attachment", "cli._run_command", "Full bug-filing flow: AI structure → pickers → preview → create"],
        ["commands/task_create.py",     "httpx, questionary, rich, sys, re\nconfig, jira_client, ai, models\ncommands/_helpers, commands/_attachment", "cli._run_command", "Full task-create flow: epic validation → AI description → create → Done"],
        ["commands/mk_bugsheet.py",     "httpx, rich, sys, pathlib\nconfig, jira_client, excel\ncommands/_helpers", "cli._run_command",                               "Project/epic pickers → fetch bugs → write .xlsx → print path"],
    ],
    col_widths=[1.6, 1.8, 1.8, 1.3],
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. DATA FLOW
# ═══════════════════════════════════════════════════════════════════════════
h1("4. DATA FLOW")

# ── 4.1 jira mk bug ─────────────────────────────────────────────────────────
h2("4.1  jira mk bug  — Filing an AI-structured bug")

h3("Step 1 — User invokes the command")
body("User runs jira → picks 'File a bug with AI description' from the menu.")
code("_run_command('mk', ['bug'])  →  from qa_jira.commands.mk_bug import run; run()")

h3("Step 2 — Collect raw description + environment")
body("commands/mk_bug.py  —  run()")
code(
    "raw_description = questionary.text('Describe the bug:').ask()\n"
    "environment     = questionary.select('Environment:', choices=['Production','Demo','Test']).ask()"
)
body("Data at this point:")
code('raw_description = "Login page auto-logs out when app is used multiple times"\nenvironment     = "Demo"')

h3("Step 3 — Detect if screenshot is attached")
code(
    "# In commands/mk_bug.py\n"
    "image_paths: list[str] = []\n"
    "if attachment and attachment.type == 'file' and attachment.filePath:\n"
    "    if Path(attachment.filePath).suffix.lower() in IMAGE_MIME:\n"
    "        image_paths.append(attachment.filePath)"
)

h3("Step 4 — AI structures the description")
body("ai/__init__.py  —  generate_bug_description(config, raw_description, environment, attachment, image_paths)")
code(
    "raw = provider.complete_json(\n"
    "    SYSTEM_PROMPT_BUG,\n"
    "    build_bug_user_prompt(raw_description),\n"
    "    max_tokens=4000,\n"
    "    image_paths=image_paths or [],\n"
    ")"
)
body("ai/http_provider.py  —  complete_json() → _post()")
body("HTTP request sent:")
code(
    "POST https://openrouter.ai/api/v1/chat/completions\n"
    "Authorization: Bearer sk-or-v1-...\n"
    "Content-Type: application/json\n\n"
    '{\n'
    '  "model": "nvidia/nemotron-3-nano-30b-a3b:free",\n'
    '  "max_tokens": 4000,\n'
    '  "messages": [\n'
    '    {"role": "system", "content": "You are a QA engineer writing structured bug reports..."},\n'
    '    {"role": "user",   "content": "Convert this bug description..."}\n'
    '  ]\n'
    '}'
)
body("If model returns 404 (offline), the provider silently tries the next model in the fallback chain.")
body("AI response (raw JSON string from model):")
code(
    '{\n'
    '  "title": "Login page auto-logs out during repeated app usage",\n'
    '  "stepsToReproduce": [\n'
    '    "Open the Heritage app on Demo environment",\n'
    '    "Log in with supervisor credentials (e.g. suresht)",\n'
    '    "Navigate to any main feature (e.g. Work With)",\n'
    '    "Return to home and navigate again 2-3 times",\n'
    '    "Observe the session state"\n'
    '  ],\n'
    '  "actualResult": "The user is unexpectedly logged out after ...",\n'
    '  "expectedResult": "The session should persist across navigation ...",\n'
    '  "additionalContext": "Reproduced on Demo environment with supervisor role ..."\n'
    '}'
)

h3("Step 5 — Parse JSON + build ADF")
body("ai/base.py  —  parse_json_loose(raw) → build_bug_result(parsed, raw_description, environment)")
code(
    "def build_bug_result(parsed, raw_description, environment, attachment=None):\n"
    "    title  = parsed.get('title') or raw_description[:80]\n"
    "    steps  = parsed.get('stepsToReproduce') or ['Reproduce using description']\n"
    "    actual = parsed.get('actualResult') or raw_description\n"
    "    # ... builds ADF blocks ...\n"
    "    return AIBugResult(\n"
    "        title=title, stepsToReproduce=steps,\n"
    "        actualResult=actual, expectedResult=expected,\n"
    "        adf=make_doc(blocks), preview=preview_string\n"
    "    )"
)
body("ADF (Atlassian Document Format) shape passed to Jira:")
code(
    '{\n'
    '  "type": "doc", "version": 1,\n'
    '  "content": [\n'
    '    {"type": "paragraph", "content": [{"type": "text", "text": "Steps to Reproduce", "marks": [{"type": "strong"}]}]},\n'
    '    {"type": "bulletList", "content": [ {"type": "listItem", "content": [...]} ]},\n'
    '    ...\n'
    '  ]\n'
    '}'
)

h3("Step 6 — Project / epic / assignee pickers")
body("commands/_helpers.py  —  pick_project(), pick_epic(), pick_user()")
body("Each picker hits the Jira REST API:")
code(
    "# search_projects\n"
    "GET https://applicate.atlassian.net/rest/api/3/project/search?query=hfc&maxResults=10\n"
    "Authorization: Basic base64(email:token)\n\n"
    "# search_epics_in_project (JQL)\n"
    'POST /rest/api/3/search/jql\n'
    '{"jql": "project=\\"HFC\\" AND issuetype=Epic ORDER BY created DESC", "maxResults": 10}'
)

h3("Step 7 — Auto-discover required custom fields")
body("commands/_helpers.py  —  prompt_for_required_extra_fields(client, base_url, auth, project_key, 'Bug')")
code(
    "GET /rest/api/3/issue/createmeta?projectKeys=HFC&issuetypeNames=Bug&expand=projects.issuetypes.fields\n\n"
    "# Response parsed → any required field NOT already handled is prompted:\n"
    '# "Complexity" → questionary.select(["1","2","3","4"]).ask()\n'
    '# Result: {"customfield_XXXXX": {"value": "2"}}'
)

h3("Step 8 — Create bug in Jira")
body("jira_client.py  —  create_bug()")
body("Payload sent to Jira:")
code(
    "POST /rest/api/3/issue\n"
    '{\n'
    '  "fields": {\n'
    '    "project": {"key": "HFC"},\n'
    '    "issuetype": {"name": "Bug"},\n'
    '    "summary": "Login page auto-logs out during repeated app usage",\n'
    '    "description": { ...ADF doc... },\n'
    '    "priority": {"name": "P2"},\n'
    '    "customfield_10148": {"value": "Demo"},\n'
    '    "customfield_XXXXX": {"value": "2"},\n'
    '    "parent": {"key": "HFC-315"}\n'
    '  }\n'
    '}'
)
body("Response:")
code('{"id": "10042", "key": "HFC-387", "self": "https://applicate.atlassian.net/rest/api/3/issue/10042"}')

h3("Step 9 — Upload attachment (if any)")
body("commands/_attachment.py  —  upload_attachment()")
code(
    "# File attachment:\n"
    "POST /rest/api/3/issue/HFC-387/attachments\n"
    "X-Atlassian-Token: no-check\n"
    "Content-Type: multipart/form-data\n"
    "Body: file=<binary content>\n\n"
    "# Google Sheet URL (embedded in ADF description already — no upload needed)"
)

h3("Step 10 — Transition to In Progress")
body("jira_client.py  —  transition_to_in_progress()")
code(
    "GET  /rest/api/3/issue/HFC-387/transitions\n"
    "# Finds transition with 'progress' or 'start' in name\n"
    "POST /rest/api/3/issue/HFC-387/transitions\n"
    '{"transition": {"id": "21"}}'
)

h3("Final output to user")
code(
    "══════════════════════════════════════════════════\n"
    "  🐛 Bug Created: HFC-387\n"
    "  🔗 https://applicate.atlassian.net/browse/HFC-387\n"
    "══════════════════════════════════════════════════\n"
    "  Done in 12.3s"
)
page_break()

# ── 4.2 jira task create ─────────────────────────────────────────────────────
h2("4.2  jira task create  — Creating a QA task")

h3("Step 1 — Epic validation")
body("commands/task_create.py  —  run()")
code(
    "# User types HFC-315 or pastes https://applicate.atlassian.net/browse/HFC-315\n"
    "url_match = re.search(r'/browse/([A-Za-z][A-Za-z0-9]*-\\d+)', epic_key_raw)\n"
    "if url_match:\n"
    "    epic_key_raw = url_match.group(1).upper()  # → 'HFC-315'\n\n"
    "epic = get_epic_info(client, cfg.jiraBaseUrl, auth, 'HFC-315')\n"
    "# Validates issuetype == 'Epic'; raises ValueError otherwise"
)

h3("Step 2 — Story / bug fetch (non-fatal)")
code(
    "story = fetch_issue_details(client, cfg.jiraBaseUrl, auth, story_input)\n"
    "# GET /rest/api/3/issue/HFC-293?fields=summary,description,issuetype,status\n"
    "# Returns Issue(key='HFC-293', summary='Work with', descriptionText='...', ...)"
)

h3("Step 3 — Vision: download story images")
code(
    "story_imgs = fetch_story_images(client, cfg.jiraBaseUrl, auth, story.key)\n"
    "# GET /rest/api/3/issue/HFC-293?fields=attachment\n"
    "# For each attachment with mimeType starting 'image/':\n"
    "#   GET attachment.content  →  write to /tmp/xxxxx.png\n"
    "# Returns: ['/tmp/img1.png', '/tmp/img2.png', ...]"
)

h3("Step 4 — AI generates task description")
body("Same provider/fallback chain as mk bug. Different prompts:")
code(
    "build_task_user_prompt('tested', story, bug_list, user_notes, attachment)\n"
    "# Returns prompt demanding:\n"
    "#   summary  : 5-7 sentences about what was tested (min 100 words)\n"
    "#   details  : 5-7 sentences about testing approach (min 100 words)\n"
    "#   outcome  : 3-4 sentences on pass/fail status (min 60 words)"
)

h3("Step 5 — Create task + mark Done")
code(
    "POST /rest/api/3/issue\n"
    '{\n'
    '  "fields": {\n'
    '    "project": {"key": "HFC"},\n'
    '    "parent":  {"key": "HFC-315"},\n'
    '    "issuetype": {"name": "Task"},\n'
    '    "summary": "QA Testing — Work With",\n'
    '    "description": { ...ADF... },\n'
    '    "assignee": {"accountId": "..."},\n'
    '    "duedate": "2026-06-01",\n'
    '    "customfield_10015": "2026-06-01",\n'
    '    "customfield_XXXXX": {"value": "2"}  # Complexity\n'
    '  }\n'
    '}\n\n'
    "# Then: transition to Done\n"
    "POST /rest/api/3/issue/HFC-389/transitions  {transition: {id: '31'}}"
)
page_break()

# ── 4.3 jira mk bugsheet ────────────────────────────────────────────────────
h2("4.3  jira mk bugsheet  — Exporting bugs to Excel")

h3("Step 1 — Pick project + epic")
body("commands/mk_bugsheet.py  →  commands/_helpers.py  →  jira_client.search_projects / search_epics_in_project")
body("Epic can be picked by: partial name search (JQL), direct key (HFC-27), or browse URL.")

h3("Step 2 — Fetch all bugs in epic")
body("jira_client.py  —  fetch_bugs_in_epic()")
code(
    "# Tries 3 JQL variants in order until one succeeds:\n"
    "jql_options = [\n"
    '    f\'"Epic Link" = {epic_key} AND issuetype = Bug ORDER BY created ASC\',\n'
    '    f\'cf[10014] = {epic_key} AND issuetype = Bug ORDER BY created ASC\',\n'
    '    f\'parent = {epic_key} AND issuetype = Bug ORDER BY created ASC\',\n'
    "]\n\n"
    "# Fields fetched:\n"
    "['summary','status','priority','assignee','reporter',\n"
    " 'created','description','issuetype','environment','customfield_10148']\n\n"
    "# Environment priority order:\n"
    "env = (\n"
    "    (fields.get('customfield_10148') or {}).get('value')  # actual Jira field\n"
    "    or fields.get('environment', '').strip()              # built-in field\n"
    "    or _extract_environment(desc_text)                    # keyword fallback\n"
    ")"
)
body("Returns list[BugInEpic] — each with: key, summary, status, priority, assignee, reporter, created, environment, url.")

h3("Step 3 — Write Excel file")
body("excel.py  —  write_bugsheet(bugs, epic, output_dir)")
code(
    "# Filename: bugsheet-HFC-315-2026-06-01.xlsx\n\n"
    "# Column layout:\n"
    "HEADERS = ['Bug ID','Bug Type','Reported By','Reporting Date',\n"
    "           'JIRA ID','Title','Current Status','Environment',\n"
    "           'Priority','RCA','Assignee','Remarks']\n\n"
    "# Column E gets a hyperlink formula:\n"
    'ws.cell(row=2, column=5, value=\'=HYPERLINK("https://applicate.atlassian.net/browse/HFC-99","HFC-99")\')\n\n'
    "# Styling:\n"
    "HEADER_FILL = PatternFill('solid', fgColor='1565C0')  # dark blue\n"
    "HEADER_FONT = Font(bold=True, color='FFFFFF')\n"
    "ROW_FILL_LIGHT = PatternFill('solid', fgColor='E3F2FD')  # alternating light blue\n"
    "ws.freeze_panes = 'A2'  # frozen header row"
)

h3("Final output")
code(
    "══════════════════════════════════════════════════════════\n"
    "  ✅ Bug Sheet Created\n"
    "  Epic:  HFC-315 — QA TASK\n"
    "  Bugs:  23 bugs exported\n"
    "  File:  /Users/salescode/bugsheet-HFC-315-2026-06-01.xlsx\n"
    "══════════════════════════════════════════════════════════"
)
page_break()

# ── 4.4 jira setup ───────────────────────────────────────────────────────────
h2("4.4  jira setup  — First-time configuration")

h3("Step 1 — Jira credentials")
code(
    "GET https://applicate.atlassian.net/rest/api/3/myself\n"
    "Authorization: Basic base64(email:token)\n\n"
    "Response → accountId, displayName"
)

h3("Step 2 — AI provider selection + key validation")
code(
    "# Validation call (same for all providers):\n"
    "POST {base_url}/chat/completions\n"
    'Authorization: Bearer {api_key}\n'
    '{"model": "...", "max_tokens": 10, "messages": [{"role":"user","content":"Say ok"}]}'
)

h3("Step 3 — Save config")
code(
    "# Written to ~/.qa-jira/config.json  (chmod 600)\n"
    '{\n'
    '  "jiraEmail":    "hiya.arora@salescode.ai",\n'
    '  "jiraApiToken": "ATATT...",\n'
    '  "jiraBaseUrl":  "https://applicate.atlassian.net",\n'
    '  "accountId":    "712020:...",\n'
    '  "displayName":  "Hiya Arora",\n'
    '  "aiProvider":   "openrouter",\n'
    '  "aiApiKey":     "sk-or-v1-...",\n'
    '  "aiModel":      "nvidia/nemotron-3-nano-30b-a3b:free",\n'
    '  "aiBaseUrl":    "https://openrouter.ai/api/v1"\n'
    '}'
)
page_break()

# ── 4.5 jira rm ──────────────────────────────────────────────────────────────
h2("4.5  jira rm  — Deleting a Jira issue")

h3("Step 1 — Fetch + display issue")
code(
    "key = extract_issue_key('HFC-387')  # also accepts browse URLs\n"
    "GET /rest/api/3/issue/HFC-387?fields=summary,description,issuetype,status,priority,assignee"
)

h3("Step 2 — Confirm (default No)")
code(
    'confirmed = questionary.confirm(\'Delete HFC-387 — "Login page auto-logs out"?\', default=False).ask()\n'
    "if not confirmed:\n"
    "    sys.exit(0)"
)

h3("Step 3 — Delete")
code(
    "DELETE /rest/api/3/issue/HFC-387\n"
    "Authorization: Basic ...\n\n"
    "# 403 → 'Permission denied' message with tip\n"
    "# 404 → 'Issue not found — may have already been deleted'"
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. ARCHITECTURE DIAGRAM
# ═══════════════════════════════════════════════════════════════════════════
h1("5. ARCHITECTURE DIAGRAM")
body("Full ASCII diagram showing all layers, files, and external services.")
doc.add_paragraph()
code(r"""
 ╔══════════════════════════════════════════════════════════════════════╗
 ║                        USER TERMINAL                                ║
 ╚═══════════════════════════════╤══════════════════════════════════════╝
                                 │  types: jira  (or picks from menu)
                                 ▼
 ╔═══════════════════════════════════════════════════════════════════════╗
 ║  ENTRY POINT  (~/.local/bin/jira → pyproject.toml console_script)   ║
 ║                                                                       ║
 ║   cli.py → main()                                                     ║
 ║     ├─ no args  → _interactive_menu()  [questionary.select]           ║
 ║     └─ with args → _run_command(cmd, rest)                            ║
 ╚══════════════════════════╤════════════════════════════════════════════╝
                            │  dispatches to command module
        ┌───────────────────┼────────────────────────────┐
        ▼                   ▼                   ▼         ▼
 ┌─────────────┐  ┌──────────────┐  ┌─────────────┐  ┌──────┐
 │ setup.py    │  │ task_create  │  │  mk_bug.py  │  │ rm   │
 │             │  │ .py          │  │             │  │ .py  │
 └──────┬──────┘  └──────┬───────┘  └──────┬──────┘  └──┬───┘
        │                │                  │             │
        │         ┌──────▼──────────────────▼─────┐      │
        │         │     commands/_helpers.py        │      │
        │         │  pick_project()  pick_epic()    │      │
        │         │  pick_user()     ask_attachment()│     │
        │         │  prompt_for_required_extra_fields│     │
        │         └──────────────┬─────────────────┘      │
        │                        │                         │
        └──────────┐   ┌─────────┘   ┌───────────────────-┘
                   ▼   ▼             ▼
 ╔══════════════════════════════════════════════════════════════════════╗
 ║                     SHARED UTILITY LAYER                            ║
 ║                                                                      ║
 ║  config.py          models.py         adf.py                        ║
 ║  get_config()       Config            make_doc()                    ║
 ║  save_config()      Issue             make_paragraph()              ║
 ║  ~/.qa-jira/        BugInEpic         make_text()                   ║
 ║  config.json        AIBugResult       make_link()                   ║
 ║                     AITaskResult      make_bullet_list()            ║
 ║                                                                      ║
 ║  file_handler.py    excel.py                                        ║
 ║  detect_input_type  write_bugsheet()  ← openpyxl                   ║
 ║  validate_file()    .xlsx output                                    ║
 ╚══════════════════════════════════════════════════════════════════════╝
        │  (auth header, REST calls)               │  (AI calls)
        ▼                                          ▼
 ╔═══════════════════════════╗   ╔════════════════════════════════════╗
 ║   jira_client.py          ║   ║   ai/  module                      ║
 ║                           ║   ║                                    ║
 ║  basic_auth_header()      ║   ║  ai/__init__.py                    ║
 ║  fetch_issue_details()    ║   ║    generate_bug_description()      ║
 ║  get_epic_info()          ║   ║    generate_task_description()     ║
 ║  search_projects()        ║   ║                                    ║
 ║  search_epics_in_project()║   ║  ai/base.py                        ║
 ║  search_users()           ║   ║    parse_json_loose()              ║
 ║  create_task()            ║   ║    build_bug_result()              ║
 ║  create_bug()             ║   ║    build_task_result()             ║
 ║  transition_to_done()     ║   ║                                    ║
 ║  transition_to_progress() ║   ║  ai/http_provider.py               ║
 ║  delete_issue()           ║   ║    HttpProvider                    ║
 ║  attach_file_to_issue()   ║   ║    complete_json(image_paths)      ║
 ║  add_comment_with_link()  ║   ║    _fetch_free_models()            ║
 ║  fetch_bugs_in_epic()     ║   ║    fallback chain                  ║
 ║  fetch_story_images()     ║   ║    vision model detection          ║
 ║  get_required_extra_fields║   ║                                    ║
 ╚══════════╤════════════════╝   ╚════════════════╤═══════════════════╝
            │  HTTPS + Basic Auth                  │  HTTPS + Bearer Token
            ▼                                      ▼
 ╔══════════════════════════╗   ╔════════════════════════════════════╗
 ║  JIRA CLOUD              ║   ║  AI PROVIDER (OpenRouter / Anthro) ║
 ║  applicate.atlassian.net ║   ║                                    ║
 ║                          ║   ║  GET /api/v1/models                ║
 ║  /rest/api/3/issue       ║   ║    → list of free models           ║
 ║  /rest/api/3/search/jql  ║   ║  POST /api/v1/chat/completions     ║
 ║  /rest/api/3/project     ║   ║    text-only or multimodal         ║
 ║  /rest/api/3/user        ║   ║    (base64 images in content[])    ║
 ║  /rest/api/3/transitions ║   ║  → JSON string with bug structure  ║
 ║  attachments/            ║   ║                                    ║
 ║  MongoDB: (none — Jira   ║   ╚════════════════════════════════════╝
 ║  is the data store)      ║
 ╚══════════════════════════╝
        │
        ▼  (local file system)
 ╔═══════════════════════════════════╗
 ║  LOCAL OUTPUTS                    ║
 ║  ~/.qa-jira/config.json (chmod 600║
 ║  bugsheet-HFC-315-2026-06-01.xlsx ║
 ╚═══════════════════════════════════╝
""")
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. KEY DESIGN DECISIONS
# ═══════════════════════════════════════════════════════════════════════════
h1("6. KEY DESIGN DECISIONS")

h2("6.1  Raw httpx instead of Anthropic/OpenAI SDK")
body(
    "The original Node.js CLI made plain axios.post() calls to the OpenAI-compatible "
    "chat/completions endpoint using just a Bearer token. The Python port deliberately "
    "mirrors this: one HttpProvider class posts to {base_url}/chat/completions with "
    "Authorization: Bearer. This means ANY provider that speaks the OpenAI chat format "
    "works — Anthropic, OpenRouter, Groq, Ollama, or a custom endpoint — without "
    "installing multiple SDKs. Switching providers is a config-file change, not a "
    "code change."
)

h2("6.2  Live model discovery instead of hardcoded list")
body(
    "OpenRouter's free model catalogue changes daily — models go offline without warning. "
    "Instead of a static list that breaks silently, the provider calls GET "
    "/api/v1/models at the start of every AI call and filters to models where "
    "pricing.prompt == '0'. This list is prepended to the seed fallback list. "
    "The first model that returns a 200 response wins. The user never sees "
    "'model not found' errors."
)

h2("6.3  createmeta API for required-field discovery")
body(
    "Different Jira workspaces have different required custom fields (e.g. Complexity, "
    "Story Points, Sprint). Hardcoding these field IDs would break the CLI on any "
    "other workspace. Instead, before creating an issue the CLI calls "
    "/rest/api/3/issue/createmeta?expand=projects.issuetypes.fields to fetch the "
    "full field schema for that project + issue type combination. Any field marked "
    "required=true that is not already handled by the CLI is presented to the user "
    "as an interactive prompt with allowed values. The result is merged into the "
    "fields payload. This makes the CLI workspace-agnostic."
)

h2("6.4  Pydantic v2 for all data models")
body(
    "Every object that crosses a boundary — config file, Jira API response, AI output, "
    "attachment metadata — is represented as a Pydantic v2 model. This gives "
    "field-level validation with clear error messages, .model_dump() for JSON "
    "serialisation, and model_validate() for deserialisation. It also makes the "
    "codebase self-documenting: any reader can see the exact shape of each data object "
    "without reading the API docs."
)

h2("6.5  Synchronous httpx throughout")
body(
    "The CLI is a short-lived interactive process — one command, then exit. "
    "There is no benefit to async I/O because requests happen sequentially "
    "(search project → search epic → create issue → upload attachment). "
    "Using synchronous httpx.Client keeps the code simple: no event loop, "
    "no await keywords, no asyncio.run() wrapper. All HTTP calls are wrapped "
    "in a single with httpx.Client(timeout=30) as client: block per command "
    "to ensure the connection is cleanly closed."
)

h2("6.6  ADF (Atlassian Document Format) builder")
body(
    "Jira's REST API v3 requires issue descriptions in ADF — a JSON tree of "
    "paragraph, bulletList, text, and link nodes. adf.py provides a minimal "
    "set of builder functions (make_doc, make_paragraph, make_text, make_link, "
    "make_bullet_list) that assemble this tree. This keeps the JSON structure "
    "out of the AI and command modules, and makes it easy to add new node types "
    "(e.g. heading, codeBlock) without touching any command code."
)

h2("6.7  Config stored as plain JSON with chmod 600")
body(
    "Credentials are stored at ~/.qa-jira/config.json with OS-level read "
    "protection (chmod 600 = owner read/write only). This mirrors how SSH keys "
    "and git credentials are conventionally stored on Unix systems. No keychain "
    "integration is needed for a developer tool, and the plain JSON format means "
    "any change to the schema is trivially backward-compatible."
)

h2("6.8  Vision AI: base64 in-message encoding")
body(
    "When the user attaches a screenshot or when the referenced Jira story has "
    "image attachments, the CLI encodes them as base64 data URLs and includes "
    "them in the user message content array — the standard OpenAI multimodal "
    "format: [{type: text, text: prompt}, {type: image_url, image_url: {url: "
    "data:image/png;base64,...}}]. A separate list of vision-capable model IDs "
    "is tried first; if none are available the call falls back to text-only models "
    "with the images stripped. Temp image files downloaded from Jira are deleted "
    "after the AI call completes."
)

h2("6.9  Google Sheet URLs embedded in ADF, not uploaded as files")
body(
    "Jira attachments must be binary files — a URL cannot be uploaded as a Jira "
    "attachment directly. Early versions of the CLI added the URL as a comment, "
    "which was easy to miss. The current design asks the user for a short label "
    "('What is this Google Sheet? e.g. Test Cases') and embeds a bold label + "
    "clickable hyperlink directly in the issue description body using ADF link "
    "nodes. This makes the sheet immediately visible to anyone reading the ticket."
)

h2("6.10  Where to make changes")
body("Quick guide for when you want to modify or extend the project:")
doc.add_paragraph()
table(
    ["If you want to...", "File(s) to change"],
    [
        ["Add a new command",                     "cli.py (_MENU list + _run_command), create commands/new_cmd.py"],
        ["Change AI prompt for bug descriptions", "prompts.py (build_bug_user_prompt) + ai/base.py (build_bug_result for ADF shape)"],
        ["Change AI prompt for task descriptions","prompts.py (build_task_user_prompt) + ai/base.py (build_task_result)"],
        ["Add a new AI provider",                 "ai/http_provider.py (add base URL constant), commands/setup.py (PROVIDER_PRESETS)"],
        ["Add a new required field handler",      "commands/_helpers.py (prompt_for_required_extra_fields) + jira_client.py (_HANDLED_FIELD_IDS)"],
        ["Change Excel column layout/styling",    "excel.py (HEADERS list, write_bugsheet function)"],
        ["Add a new Jira API call",               "jira_client.py (add function), models.py (add return type if needed)"],
        ["Change config fields",                  "models.py (Config class), config.py (load/save), commands/setup.py (wizard questions)"],
        ["Add a new attachment type",             "file_handler.py (EXTENSION_LABELS, detect_input_type), commands/_attachment.py"],
        ["Swap Jira workspace URL",               "commands/setup.py (JIRA_BASE_URL constant)"],
        ["Add a new Pydantic model",              "models.py"],
        ["Add or update a test",                  "tests/test_*.py — use httpx.MockTransport for Jira calls, unittest.mock for AI"],
    ],
    col_widths=[3.0, 3.5],
)

page_break()
h1("7. QUICK REFERENCE — All CLI Commands")
table(
    ["Command", "What it does", "Key files involved"],
    [
        ["jira",                "Opens interactive arrow-key menu",              "cli.py"],
        ["jira setup",          "4-step config wizard (Jira + AI)",              "commands/setup.py, config.py, jira_client.py"],
        ["jira task create",    "AI-structured QA task under an epic",           "commands/task_create.py, ai/*, jira_client.py"],
        ["jira mk bug",         "AI-structured bug report",                      "commands/mk_bug.py, ai/*, jira_client.py"],
        ["jira mk bugsheet",    "Export all epic bugs to .xlsx",                 "commands/mk_bugsheet.py, excel.py, jira_client.py"],
        ["jira rm <KEY|URL>",   "Delete a Jira issue (confirm-No-by-default)",   "commands/rm.py, jira_client.py"],
    ],
    col_widths=[1.6, 2.8, 2.1],
)

doc.save(OUT_PATH)
print(f"Saved → {OUT_PATH}")
